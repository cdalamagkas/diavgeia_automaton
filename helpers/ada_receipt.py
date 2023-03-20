import docx.oxml
import docx.opc.constants
from docx import Document
import requests
from docx.shared import Pt


def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )
    hyperlink.set(docx.oxml.shared.qn('w:history'), '1', )  #CDAL

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    rStyle = docx.oxml.shared.OxmlElement('w:rStyle')       #CDAL
    rStyle.set(docx.oxml.shared.qn('w:val'), 'Hyperlink', ) #CDAL

    # Join all the xml elements together add add the required text to the w:r element
    rPr.append(rStyle)  #CDAL
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    #r = paragraph.add_run()
    #r._r.append(hyperlink)

    paragraph._p.append(hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    #r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    #r.font.underline = True

    return hyperlink
    

def create_ada_receipt(ada, CONFIG):
    try:
        ada_payload = requests.get(CONFIG["BASE_URL"] + 'decisions/' + ada, headers={'Accept': 'application/json'})
        ada_payload = ada_payload.json()

        if 'errors' in ada_payload:
            print('[!!] Ο ΑΔΑ που δώσατε δεν είναι έγκυρος ή υπάρχει πρόβλημα επικοινωνίας με τη Δι@ύγεια. Παρακαλώ προσπαθήστε ξανά.')
            print('=============================================')
            return False

    except requests.ConnectionError:
        print("[!!] Σφάλμα σύνδεσης με τη Δι@ύγεια.")
        print("====================================")
        return False
    
    filename = ada_payload['protocolNumber'] + '-' + ada + '.docx'
    
    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.size = Pt(10)
    style.paragraph_format.space_after = Pt(0)

    p = document.add_paragraph()
    run = p.add_run('Η ανάρτηση της πράξης ολοκληρώθηκε με επιτυχία.')
    run.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(6)

    p = document.add_paragraph()
    run = p.add_run('ΑΔΑ:')
    run.font.bold = True
    p.add_run(' ')
    p.add_run(ada)

    p = document.add_paragraph()
    run = p.add_run('Φορέας έκδοσης:')
    run.font.bold = True
    p.add_run(' ')
    p.add_run(CONFIG["FOREAS_EKDOSIS"])

    p = document.add_paragraph()
    run = p.add_run('Αριθμός πρωτοκόλλου:')
    run.font.bold = True
    p.add_run(' ')
    p.add_run(ada_payload['protocolNumber'])

    p = document.add_paragraph()
    run = p.add_run('Θέμα:')
    run.font.bold = True
    p.add_run(' ')
    p.add_run(ada_payload['subject'])

    p = document.add_paragraph()
    run = p.add_run('Είδος Πράξης:')
    run.font.bold = True
    p.add_run(' ')
    p.add_run('ΑΝΑΛΗΨΗ ΥΠΟΧΡΕΩΣΗΣ')

    p = document.add_paragraph(' ')

    p = document.add_paragraph('Για προβολή της αναρτημένης πράξης ')
    add_hyperlink(p, 'πατήστε εδώ', 'https://diavgeia.gov.gr/decision/view/'+ada)
    p = document.add_paragraph('Για λήψη του εγγράφου της πράξης ')
    add_hyperlink(p, 'πατήστε εδώ', 'https://diavgeia.gov.gr/doc/'+ada)

    run = document.add_paragraph().add_run('* Για την ορθή προβολή του εγγράφου και την επαλήθευση της ψηφιακής υπογραφής προτείνεται η χρήση Adobe Reader.')
    run.font.size = Pt(8.5)

    document.save(CONFIG["RECEIPTS_DIR"] + "/" + filename)
    
    print("[✓] Η απόδειξη με όνομα " + filename + " δημιουργήθηκε.")
    
    return True
